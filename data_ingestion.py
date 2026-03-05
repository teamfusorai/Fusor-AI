import os
import tempfile
import uuid
from typing import Optional, List

from fastapi import Request, UploadFile, File, Form, APIRouter
from pydantic import BaseModel
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, PointStruct, Filter, FieldCondition, MatchValue

import PyPDF2
import docx
from utils.sitemap import get_sitemap_urls
from utils.logging_config import get_logger
from utils.metrics import increment
import config

# Set environment variable to disable symlinks on Windows
os.environ["HF_HUB_DISABLE_SYMLINKS_WARNING"] = "1"

# -------------------------
# Setup
# -------------------------
router = APIRouter()
load_dotenv()
logger = get_logger(__name__)

# Shared result types for ingestion (file or URL)
class _SimpleDocument:
    def __init__(self, text: str):
        self.text = text
    def export_to_markdown(self) -> str:
        return self.text

class _SimpleResult:
    def __init__(self, text: str, url: Optional[str] = None):
        self.document = _SimpleDocument(text)
        self.url = url

# Initialize docling for URL processing, fallback for documents
try:
    from docling.document_converter import DocumentConverter
    converter = DocumentConverter()
    DOCLING_AVAILABLE = True
    logger.info("Docling available for URL processing")
except ImportError:
    converter = None
    DOCLING_AVAILABLE = False
    logger.info("Docling not available, using fallback for all processing")

def extract_text_fallback(file_path: str) -> str:
    """Fallback text extraction for when Docling fails"""
    try:
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
                
        elif file_ext in ['.docx', '.doc']:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
            
        elif file_ext in ['.txt', '.md']:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
                
        else:
            return f"Unsupported file type: {file_ext}"
            
    except Exception as e:
        return f"Error extracting text: {str(e)}"

def extract_text_from_url(url: str) -> str:
    """Extract text from URL using requests and BeautifulSoup"""
    try:
        import requests
        from bs4 import BeautifulSoup
        
        # Add headers to mimic a real browser
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Accept-Encoding': 'gzip, deflate',
            'Connection': 'keep-alive',
        }
        
        response = requests.get(url, timeout=15, headers=headers)
        response.raise_for_status()
        
        # Parse HTML content
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()
        
        # Extract text
        text = soup.get_text()
        
        # Clean up text
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        return text.strip()
        
    except Exception as e:
        return f"Error extracting text from URL: {str(e)}"

QDRANT_HOST = config.QDRANT_HOST
QDRANT_PORT = config.QDRANT_PORT

EMBEDDING_MODEL = OpenAIEmbeddings(
    model=config.EMBEDDING_MODEL_NAME,
    api_key=os.getenv("OPENAI_API_KEY"),
    timeout=config.TIMEOUT_SECONDS,
    max_retries=config.MAX_RETRIES,
)

qdrant_client = QdrantClient(host=QDRANT_HOST, port=QDRANT_PORT)

class IngestRequest(BaseModel):
    url: Optional[str] = None
    user_id: Optional[str] = None
    bot_id: Optional[str] = None


def run_ingest_sync(
    source_type: str,
    path_or_url: str,
    source_name: str,
    user_id: str,
    bot_id: str,
) -> dict:
    """
    Run ingestion synchronously (for Celery worker).
    source_type: "file" | "url"
    path_or_url: file path or URL string
    Returns dict with status, chunks_stored, error, etc.
    """
    user_id = user_id or "default_user"
    bot_id = bot_id or "default_bot"
    collection_name = f"{config.COLLECTION_PREFIX}{user_id}"
    namespace = f"{config.NAMESPACE_PREFIX}{bot_id}"
    sources: List[object] = []

    if source_type == "file":
        text = extract_text_fallback(path_or_url)
        if not text or text.startswith("Error"):
            return {"status": "error", "error": f"Failed to extract text: {text}"}
        sources.append(_SimpleResult(text))
    elif source_type == "url":
        sitemap_base = path_or_url.replace("/sitemap.xml", "").replace("sitemap.xml", "").rstrip("/") or path_or_url
        if DOCLING_AVAILABLE and converter:
            if "sitemap.xml" in path_or_url:
                try:
                    sitemap_urls = get_sitemap_urls(sitemap_base)
                    sources = list(converter.convert_all(sitemap_urls))
                except Exception as e:
                    logger.warning("Sitemap failed, single page", extra={"error": str(e)})
                    sources = [converter.convert(path_or_url)]
            else:
                sources = [converter.convert(path_or_url)]
        else:
            if "sitemap.xml" in path_or_url:
                try:
                    sitemap_urls = get_sitemap_urls(sitemap_base)
                    for u in sitemap_urls:
                        t = extract_text_from_url(u)
                        if t and not t.startswith("Error"):
                            sources.append(_SimpleResult(t, u))
                except Exception:
                    t = extract_text_from_url(path_or_url)
                    if t and not t.startswith("Error"):
                        sources.append(_SimpleResult(t, path_or_url))
            else:
                t = extract_text_from_url(path_or_url)
                if t and not t.startswith("Error"):
                    sources.append(_SimpleResult(t, path_or_url))
    else:
        return {"status": "error", "error": "Invalid source_type"}

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.CHUNK_SIZE,
        chunk_overlap=config.CHUNK_OVERLAP,
    )
    all_chunks = []
    for result in sources:
        if not getattr(result, "document", None):
            continue
        text = result.document.export_to_markdown()
        chunks = splitter.create_documents([text])
        for i, c in enumerate(chunks):
            c.metadata.update({
                "namespace": namespace,
                "source_name": source_name,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "user_id": user_id,
                "bot_id": bot_id,
            })
        all_chunks.extend(chunks)

    if not all_chunks:
        return {"status": "error", "error": "No content to embed."}

    if not qdrant_client.collection_exists(collection_name):
        qdrant_client.create_collection(
            collection_name=collection_name,
            vectors_config=VectorParams(
                size=config.EMBEDDING_DIMENSIONS,
                distance=getattr(Distance, config.EMBEDDING_DISTANCE_METRIC),
            ),
        )
    try:
        qdrant_client.delete(
            collection_name=collection_name,
            points_selector=Filter(
                must=[
                    FieldCondition(key="metadata.namespace", match=MatchValue(value=namespace)),
                    FieldCondition(key="metadata.source_name", match=MatchValue(value=source_name)),
                ]
            ),
        )
    except Exception:
        pass

    texts = [c.page_content for c in all_chunks]
    batch_size = config.BATCH_SIZE
    all_vectors = []
    for i in range(0, len(texts), batch_size):
        batch = texts[i : i + batch_size]
        vectors = EMBEDDING_MODEL.embed_documents(batch)
        all_vectors.extend(vectors)
    points = [
        PointStruct(
            id=str(uuid.uuid4()),
            vector=vec,
            payload={"page_content": doc.page_content, "metadata": doc.metadata},
        )
        for doc, vec in zip(all_chunks, all_vectors)
    ]
    for i in range(0, len(points), batch_size):
        qdrant_client.upsert(collection_name=collection_name, points=points[i : i + batch_size])

    logger.info("Ingest completed (sync)", extra={"chunks_stored": len(all_chunks), "collection": collection_name})
    return {
        "status": "success",
        "chunks_stored": len(all_chunks),
        "qdrant_collection": collection_name,
        "namespace": namespace,
        "user_id": user_id,
        "bot_id": bot_id,
        "source_name": source_name,
    }


def _ensure_upload_dir() -> None:
    d = config.INGEST_UPLOAD_DIR
    if d and not os.path.isdir(d):
        os.makedirs(d, exist_ok=True)


# -------------------------
# Ingestion Endpoint
# -------------------------
@router.post("/ingest")
async def ingest(
    request: Request,
    file: Optional[UploadFile] = File(None),
    url: Optional[str] = Form(None),
    user_id: Optional[str] = Form(None),
    bot_id: Optional[str] = Form(None)
):
    increment("ingest_requests_total")
    uid = user_id or "default_user"
    bid = bot_id or "default_bot"
    source_name = "unknown"

    if not file and not url:
        increment("ingest_errors_total")
        return {"error": "Either 'file' or 'url' must be provided."}

    # When Celery is configured, enqueue and return job_id
    if config.CELERY_BROKER_URL:
        try:
            from tasks.ingest_tasks import run_ingest as run_ingest_task
        except ImportError:
            logger.warning("Celery broker set but tasks not importable; running inline")
        else:
            if file:
                content = await file.read()
                if len(content) > config.MAX_INGEST_FILE_BYTES:
                    increment("ingest_errors_total")
                    return {"error": f"File too large. Maximum size: {config.MAX_INGEST_FILE_BYTES} bytes."}
                source_name = file.filename or "uploaded_file"
                _ensure_upload_dir()
                unique_id = str(uuid.uuid4())
                suffix = os.path.splitext(file.filename or "")[1] or ""
                path = os.path.join(config.INGEST_UPLOAD_DIR, f"{unique_id}{suffix}")
                with open(path, "wb") as f:
                    f.write(content)
                result = run_ingest_task.delay("file", path, source_name, uid, bid)
                return {"status": "queued", "job_id": result.id, "message": "Ingestion queued; poll GET /ingest/status/{job_id}"}
            else:
                source_name = url
                result = run_ingest_task.delay("url", url, source_name, uid, bid)
                return {"status": "queued", "job_id": result.id, "message": "Ingestion queued; poll GET /ingest/status/{job_id}"}

    # Inline path (no Celery)
    sources = []
    if file:
        content = await file.read()
        if len(content) > config.MAX_INGEST_FILE_BYTES:
            increment("ingest_errors_total")
            return {"error": f"File too large. Maximum size: {config.MAX_INGEST_FILE_BYTES} bytes."}
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename or "")[1]) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        try:
            text = extract_text_fallback(tmp_path)
            if text and not text.startswith("Error"):
                sources.append(_SimpleResult(text))
            else:
                increment("ingest_errors_total")
                return {"error": f"Failed to extract text: {text}"}
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
        source_name = file.filename or "uploaded_file"
    else:
        sitemap_base = url.replace("/sitemap.xml", "").replace("sitemap.xml", "").rstrip("/") or url
        if DOCLING_AVAILABLE and converter:
            if "sitemap.xml" in url:
                try:
                    sitemap_urls = get_sitemap_urls(sitemap_base)
                    sources = list(converter.convert_all(sitemap_urls))
                except Exception as e:
                    logger.warning("Sitemap failed, single page", extra={"error": str(e)})
                    sources = [converter.convert(url)]
            else:
                sources = [converter.convert(url)]
        else:
            if "sitemap.xml" in url:
                try:
                    sitemap_urls = get_sitemap_urls(sitemap_base)
                    for sitemap_url in sitemap_urls:
                        text = extract_text_from_url(sitemap_url)
                        if text and not text.startswith("Error"):
                            sources.append(_SimpleResult(text, sitemap_url))
                except Exception as e:
                    logger.warning("Sitemap failed, single page", extra={"error": str(e)})
                    text = extract_text_from_url(url)
                    if text and not text.startswith("Error"):
                        sources.append(_SimpleResult(text, url))
            else:
                text = extract_text_from_url(url)
                if text and not text.startswith("Error"):
                    sources.append(_SimpleResult(text, url))
        source_name = url

    user_id = uid
    bot_id = bid
    collection_name = f"{config.COLLECTION_PREFIX}{user_id}"
    namespace = f"{config.NAMESPACE_PREFIX}{bot_id}"
    all_chunks = []
    splitter = RecursiveCharacterTextSplitter(chunk_size=config.CHUNK_SIZE, chunk_overlap=config.CHUNK_OVERLAP)
    if not qdrant_client.collection_exists(collection_name):
        qdrant_client.create_collection(
            collection_name=collection_name,
            vectors_config=VectorParams(size=config.EMBEDDING_DIMENSIONS, distance=getattr(Distance, config.EMBEDDING_DISTANCE_METRIC)),
        )
    for result in sources:
        if not getattr(result, "document", None):
            continue
        text = result.document.export_to_markdown()
        chunks = splitter.create_documents([text])
        for i, chunk in enumerate(chunks):
            chunk.metadata.update({
                "namespace": namespace,
                "source_name": source_name,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "user_id": user_id,
                "bot_id": bot_id,
            })
        all_chunks.extend(chunks)
    if not all_chunks:
        increment("ingest_errors_total")
        return {"error": "No content to embed."}
    try:
        qdrant_client.delete(
            collection_name=collection_name,
            points_selector=Filter(
                must=[
                    FieldCondition(key="metadata.namespace", match=MatchValue(value=namespace)),
                    FieldCondition(key="metadata.source_name", match=MatchValue(value=source_name)),
                ]
            ),
        )
    except Exception:
        pass
    texts = [c.page_content for c in all_chunks]
    batch_size = config.BATCH_SIZE
    all_vectors = []
    for i in range(0, len(texts), batch_size):
        vectors = await EMBEDDING_MODEL.aembed_documents(texts[i : i + batch_size])
        all_vectors.extend(vectors)
    points = [
        PointStruct(id=str(uuid.uuid4()), vector=vec, payload={"page_content": doc.page_content, "metadata": doc.metadata})
        for doc, vec in zip(all_chunks, all_vectors)
    ]
    for i in range(0, len(points), batch_size):
        qdrant_client.upsert(collection_name=collection_name, points=points[i : i + batch_size])
    logger.info("Ingest completed", extra={"chunks_stored": len(all_chunks), "collection": collection_name})
    return {
        "status": "success",
        "chunks_stored": len(all_chunks),
        "qdrant_collection": collection_name,
        "namespace": namespace,
        "user_id": user_id,
        "bot_id": bot_id,
        "source_name": source_name,
    }


@router.get("/ingest/status/{job_id}")
async def ingest_status(job_id: str):
    """Return status and result of a queued ingestion job (requires CELERY_RESULT_BACKEND)."""
    if not config.CELERY_BROKER_URL:
        return {"error": "Background jobs not configured (no CELERY_BROKER_URL)."}
    try:
        from celery.result import AsyncResult
        from celery_app import app as celery_app
        ar = AsyncResult(job_id, app=celery_app)
        state = ar.state
        out = {"job_id": job_id, "status": state}
        if state == "SUCCESS" and ar.result:
            out["result"] = ar.result
        if state == "FAILURE" and ar.result:
            out["error"] = str(ar.result)
        return out
    except Exception as e:
        return {"error": str(e)}

# -------------------------
# Run as main
# -------------------------
if __name__ == "__main__":
    import uvicorn
    uvicorn.run("data_ingestion:router", host="0.0.0.0", port=8000, reload=True)
